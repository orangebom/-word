from bs4 import BeautifulSoup
import urllib
import re
import requests
from docx import Document
from docx.shared import Inches
#每日天文网址
astronomy_host = 'http://www.bjp.org.cn/col/col89/index.html'
picture_save = 'picture.jpg'

def get_today_picture():
    global astronomy_host
    #获取网页源代码，是字节数组
    text = urllib.request.urlopen(astronomy_host).read()
    #前一个参数为要解析的文本，后一个参数为解析模型
    # bs4的HTML解析器：BeautifulSoup(mk,'html.parser')——条件：安装bs4库
    # lxml的HTML解析器：BeautifulSoup(mk,'lxml')——pip install lxml
    # lxml的XML解析器：BeautifulSoup(mk,'xml')——pip install lxml
    # html5lib的解析器：BeautifulSoup(mk,'html5lib')——pip install html5lib
    soup = BeautifulSoup(text, 'html.parser')

    today = soup.find_all(target="_blank")
    today_host = "%s"%(today[1])
    today_host = re.findall(r"href=\"(.+?)\"",today_host);
    today_host = ''.join(today_host)

    
    return astronomy_host[0:21] + today_host


if __name__ == '__main__':

    today_host = get_today_picture()
    document = Document()
    #获取网页源代码，是字节数组
    text = urllib.request.urlopen(today_host).read()

    soup = BeautifulSoup(text, 'html.parser')
    
    #获取标题
    today_title = "%s"%(soup.title.text)
    #获取说明内容
    today_explain = soup.find_all('p')
    today_explain = today_explain[1].text
    
    #下载图片
    img = soup.find('img')
    img_down = requests.get(astronomy_host[0:21]+img['src'])
    
    #保存图片
    fw = open(picture_save,'wb')
    fw.write(img_down.content)
    fw.close()

    #编写文档
    document.add_heading(today_title, 0)  #插入标题
    document.add_picture(picture_save, width=Inches(6))
    document.add_paragraph(today_explain)   #插入段落
    
    document.save('ever_day'.docx')  #保存文档
    print(today_explain)



    